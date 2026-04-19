"""
PeerTranslate — Export Module

Converts translated Markdown into DOCX and LaTeX formats.
"""

import io
import re
import logging

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown_text: str, title: str = "PeerTranslate Output") -> bytes:
    """Convert Markdown text to a .docx file and return raw bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headers
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='Intense Quote')
            i += 1
            continue

        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        paragraph_text = line.strip()
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('>') and not lines[i].startswith('- ') and not lines[i].startswith('* '):
            paragraph_text += ' ' + lines[i].strip()
            i += 1

        # Apply bold/italic from markdown
        p = doc.add_paragraph()
        _add_formatted_text(p, paragraph_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_formatted_text(paragraph, text: str):
    """Parse markdown bold/italic and add formatted runs to docx paragraph."""
    from docx.shared import Pt
    
    # Split on bold (**text**) and italic (*text*)
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


def markdown_to_latex(markdown_text: str) -> str:
    """Convert Markdown text to LaTeX source."""
    
    preamble = r"""\documentclass[12pt, a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{lmodern}
\usepackage{amsmath, amssymb}
\usepackage{geometry}
\usepackage{hyperref}
\usepackage{graphicx}
\usepackage{booktabs}
\geometry{margin=2.5cm}
\setlength{\parindent}{0pt}
\setlength{\parskip}{0.6em}

\begin{document}

"""
    
    postamble = r"""
\end{document}
"""
    
    body = markdown_text
    
    # Escape LaTeX special chars (but not our own commands)
    body = body.replace('&', r'\&')
    body = body.replace('%', r'\%')
    body = body.replace('$', r'\$')
    body = body.replace('#', r'\#')  # we'll fix headers below
    body = body.replace('_', r'\_')
    
    # Restore headers (we broke # with escape)
    body = re.sub(r'^\\#\\#\\# (.+)$', r'\\subsubsection*{\1}', body, flags=re.MULTILINE)
    body = re.sub(r'^\\#\\# (.+)$', r'\\subsection*{\1}', body, flags=re.MULTILINE)
    body = re.sub(r'^\\# (.+)$', r'\\section*{\1}', body, flags=re.MULTILINE)
    
    # Bold and italic
    body = re.sub(r'\*\*\*(.+?)\*\*\*', r'\\textbf{\\textit{\1}}', body)
    body = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', body)
    body = re.sub(r'\*(.+?)\*', r'\\textit{\1}', body)
    
    # Bullet lists
    lines = body.split('\n')
    result_lines = []
    in_itemize = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            if not in_itemize:
                result_lines.append(r'\begin{itemize}')
                in_itemize = True
            item_text = stripped[2:]
            result_lines.append(f'  \\item {item_text}')
        else:
            if in_itemize:
                result_lines.append(r'\end{itemize}')
                in_itemize = False
            result_lines.append(line)
    if in_itemize:
        result_lines.append(r'\end{itemize}')
    
    body = '\n'.join(result_lines)
    
    # Blockquotes
    body = re.sub(r'^> (.+)$', r'\\begin{quote}\1\\end{quote}', body, flags=re.MULTILINE)
    
    # Line breaks
    body = body.replace('\n\n', '\n\n\\medskip\n\n')
    
    return preamble + body + postamble
